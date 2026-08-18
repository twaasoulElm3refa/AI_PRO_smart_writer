import json
import re
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.crud import get_existing_conversation_for_task
from app.content_common import combine_usage, normalize_text
from app.generic_content_chat import (
    calculate_content_tool_cost,
    normalize_extra_options,
    response_should_be_arabic,
)
from app.errors import ProviderOutputError
from app.json_utils import extract_json_object, object_response_format
from app.providers import send_messages_with_model
from app.schemas import (
    ChatMessage,
    ContentToolResultItem,
    GeneratedFileInfo,
    ResumeBuilderChatResponse,
    ResumeBuilderState,
)
from app.settings import get_settings
from app.tasks import RESUME_BUILDER_PROMPT

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_RESUME_EXTENSIONS = {".docx", ".txt", ".pdf"}


def _safe_filename(value: str | None, default: str = "resume") -> str:
    value = (value or default).strip()
    value = re.sub(r"[^A-Za-z0-9_.\-\u0600-\u06FF ]+", "", value)
    value = re.sub(r"\s+", "_", value).strip("._")
    return value or default


def _generated_resume_dir() -> Path:
    settings = get_settings()
    path = Path(settings.GENERATED_FILES_DIR)
    if not path.is_absolute():
        path = Path.cwd() / path
    path = path / "resumes"
    path.mkdir(parents=True, exist_ok=True)
    return path


def resume_file_path(file_id: str) -> Path:
    if not re.fullmatch(r"[a-f0-9\-]{36}", file_id):
        raise ValueError("Invalid file_id")
    return _generated_resume_dir() / f"{file_id}.docx"


def parse_resume_state(state_json: str | None) -> ResumeBuilderState:
    defaults = ResumeBuilderState(
        target_role=None,
        candidate_name=None,
        language="Auto Detect",
        tone="Professional",
        experience_level="Auto Detect",
        resume_style="ATS-friendly modern",
        output_format="docx",
        sections_to_include=["Summary", "Skills", "Experience", "Education", "Certifications", "Projects", "Languages"],
        extra_options=["Improve clarity", "Keep it honest", "ATS-friendly formatting"],
        last_output=None,
    )
    data = defaults.model_dump()
    if state_json:
        try:
            incoming = json.loads(state_json)
        except json.JSONDecodeError as exc:
            raise ValueError(f"state must be valid JSON: {exc}") from exc
        if not isinstance(incoming, dict):
            raise ValueError("state must be a JSON object")
        for key, value in incoming.items():
            if key in {"sections_to_include", "extra_options"}:
                normalized = normalize_extra_options(value)
                if normalized:
                    data[key] = normalized
            elif isinstance(value, str):
                data[key] = value.strip() or data.get(key)
            elif value is not None:
                data[key] = value
    return ResumeBuilderState(**data)


async def read_resume_upload(file: UploadFile) -> tuple[str, str]:
    filename = file.filename or "resume"
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_RESUME_EXTENSIONS:
        raise ValueError("Unsupported resume file type. Supported: .docx, .pdf, .txt")

    settings = get_settings()
    raw = await file.read()
    max_bytes = max(1, int(settings.MAX_RESUME_UPLOAD_MB)) * 1024 * 1024
    if len(raw) > max_bytes:
        raise ValueError(f"Resume file exceeds max size of {settings.MAX_RESUME_UPLOAD_MB} MB")
    if not raw:
        raise ValueError("Uploaded resume file is empty")

    if suffix == ".txt":
        for encoding in ("utf-8", "utf-8-sig", "cp1256", "latin-1"):
            try:
                return raw.decode(encoding).strip(), filename
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="ignore").strip(), filename

    if suffix == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise ValueError("python-docx is required to read DOCX resumes. Install: pip install python-docx") from exc
        doc = Document(BytesIO(raw))
        parts: list[str] = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip(), filename

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise ValueError("pypdf is required to read PDF resumes. Install: pip install pypdf") from exc
        reader = PdfReader(BytesIO(raw))
        parts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
        return "\n".join(parts).strip(), filename

    raise ValueError("Unsupported resume file type")


def build_resume_prompt(state: ResumeBuilderState, user_message: str, resume_text: str, original_filename: str) -> str:
    max_chars = 18000
    if len(resume_text) > max_chars:
        resume_text = resume_text[:max_chars].strip()
    return f"""
Latest user instruction:
{user_message}

Original uploaded filename:
{original_filename}

Resume builder state:
{json.dumps(state.model_dump(), ensure_ascii=False, indent=2)}

Extracted resume text:
{resume_text}

Task:
Create the improved resume JSON using the required JSON shape from the system prompt.

Output rules:
- Return JSON only.
- Do not include markdown.
- First character must be {{ and last character must be }}.
""".strip()


def _as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _add_bullets(doc, items: list[Any]):
    for item in items:
        text = normalize_text(item)
        if text:
            doc.add_paragraph(text, style="List Bullet")


def _add_section_heading(doc, title: str):
    doc.add_heading(title, level=2)


def create_resume_docx(resume_json: dict[str, Any], output_path: Path, state: ResumeBuilderState):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise ValueError("python-docx is required to create resume DOCX files. Install: pip install python-docx") from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    is_ar = str(state.language or "").lower() == "arabic"
    name = normalize_text(resume_json.get("candidate_name")) or normalize_text(state.candidate_name) or "Resume"
    headline = normalize_text(resume_json.get("headline")) or normalize_text(state.target_role)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_ar else WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)

    if headline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_ar else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(headline)
        r.italic = True
        r.font.size = Pt(11)

    contact = [normalize_text(x) for x in _as_list(resume_json.get("contact"))]
    contact = [x for x in contact if x]
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_ar else WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(" | ".join(contact))

    summary = normalize_text(resume_json.get("summary"))
    if summary:
        _add_section_heading(doc, "الملخص المهني" if is_ar else "Professional Summary")
        doc.add_paragraph(summary)

    skills = [normalize_text(x) for x in _as_list(resume_json.get("skills"))]
    skills = [x for x in skills if x]
    if skills:
        _add_section_heading(doc, "المهارات" if is_ar else "Skills")
        doc.add_paragraph(" • ".join(skills))

    experience = [x for x in _as_list(resume_json.get("experience")) if isinstance(x, dict)]
    if experience:
        _add_section_heading(doc, "الخبرات العملية" if is_ar else "Experience")
        for item in experience:
            role = normalize_text(item.get("role"))
            company = normalize_text(item.get("company"))
            location = normalize_text(item.get("location"))
            dates = normalize_text(item.get("dates"))
            line = " - ".join([x for x in [role, company] if x]) or company or role
            p = doc.add_paragraph()
            r = p.add_run(line or "Experience")
            r.bold = True
            details = " | ".join([x for x in [location, dates] if x])
            if details:
                p.add_run(f"\n{details}").italic = True
            _add_bullets(doc, _as_list(item.get("bullets")))

    education = [x for x in _as_list(resume_json.get("education")) if isinstance(x, dict)]
    if education:
        _add_section_heading(doc, "التعليم" if is_ar else "Education")
        for item in education:
            degree = normalize_text(item.get("degree"))
            institution = normalize_text(item.get("institution"))
            location = normalize_text(item.get("location"))
            dates = normalize_text(item.get("dates"))
            p = doc.add_paragraph()
            r = p.add_run(" - ".join([x for x in [degree, institution] if x]) or "Education")
            r.bold = True
            details = " | ".join([x for x in [location, dates] if x])
            if details:
                p.add_run(f"\n{details}").italic = True
            _add_bullets(doc, _as_list(item.get("details")))

    for key, title_en, title_ar in [
        ("certifications", "Certifications", "الشهادات"),
        ("languages", "Languages", "اللغات"),
    ]:
        items = [normalize_text(x) for x in _as_list(resume_json.get(key))]
        items = [x for x in items if x]
        if items:
            _add_section_heading(doc, title_ar if is_ar else title_en)
            _add_bullets(doc, items)

    projects = [x for x in _as_list(resume_json.get("projects")) if isinstance(x, dict)]
    if projects:
        _add_section_heading(doc, "المشاريع" if is_ar else "Projects")
        for item in projects:
            name = normalize_text(item.get("name")) or "Project"
            desc = normalize_text(item.get("description"))
            p = doc.add_paragraph()
            p.add_run(name).bold = True
            if desc:
                doc.add_paragraph(desc)
            _add_bullets(doc, _as_list(item.get("bullets")))

    additional = [x for x in _as_list(resume_json.get("additional_sections")) if isinstance(x, dict)]
    for section_data in additional:
        title = normalize_text(section_data.get("title"))
        items = _as_list(section_data.get("items"))
        if title and items:
            _add_section_heading(doc, title)
            _add_bullets(doc, items)

    doc.save(output_path)


def summarize_resume_json(resume_json: dict[str, Any]) -> str:
    lines = []
    for key in ["candidate_name", "headline", "summary"]:
        value = normalize_text(resume_json.get(key))
        if value:
            lines.append(value)
    skills = [normalize_text(x) for x in _as_list(resume_json.get("skills"))]
    skills = [x for x in skills if x]
    if skills:
        lines.append("Skills: " + ", ".join(skills[:20]))
    return "\n".join(lines).strip()


async def run_resume_builder_chat_upload(
    db: Session,
    *,
    user_id: int,
    sub_tool_id: int,
    conversation_uuid: str,
    user_message: str,
    state_json: str | None,
    file: UploadFile,
    debug: bool,
    request_id: str,
) -> ResumeBuilderChatResponse:
    settings = get_settings()

    get_existing_conversation_for_task(
        db=db,
        user_id=user_id,
        sub_tool_id=sub_tool_id,
        conversation_uuid=conversation_uuid,
    )

    user_message = (user_message or "").strip()
    if not user_message:
        raise ValueError("user_message cannot be empty")
    if len(user_message) > settings.MAX_USER_MESSAGE_LENGTH:
        raise ValueError(f"user_message exceeds max length of {settings.MAX_USER_MESSAGE_LENGTH}")

    state = parse_resume_state(state_json)
    resume_text, original_filename = await read_resume_upload(file)
    if not resume_text:
        raise ValueError("Could not extract text from uploaded resume file")

    messages = [
        ChatMessage(role="system", content=RESUME_BUILDER_PROMPT),
        ChatMessage(role="user", content=build_resume_prompt(state, user_message, resume_text, original_filename)),
    ]

    provider_result = await send_messages_with_model(
        model_key="resume_builder",
        messages=messages,
        temperature_override=0.35,
        max_tokens_override=4500,
        enable_web_search=False,
        response_format=object_response_format("resume_builder"),
    )

    try:
        resume_json = extract_json_object(provider_result.content)
    except Exception as exc:
        raise ProviderOutputError(
            "Resume builder model did not return valid JSON. "
            f"Validation error: {exc}. "
            f"Trace: {provider_result.trace_id or 'n/a'} / "
            f"{provider_result.generation_id or 'n/a'}."
        ) from exc

    candidate_name = normalize_text(resume_json.get("candidate_name")) or state.candidate_name or "resume"
    file_id = str(uuid.uuid4())
    filename = f"{_safe_filename(candidate_name)}_resume.docx"
    output_path = resume_file_path(file_id)
    try:
        create_resume_docx(resume_json, output_path, state)
    except PermissionError as exc:
        raise ValueError(
            f"Generated resume directory is not writable: {output_path.parent}. "
            "Fix folder ownership/permissions or set GENERATED_FILES_DIR to a writable path."
        ) from exc

    last_output = summarize_resume_json(resume_json)
    state.last_output = last_output or provider_result.content[:2000]
    if normalize_text(resume_json.get("candidate_name")):
        state.candidate_name = normalize_text(resume_json.get("candidate_name"))

    usage = combine_usage(provider_result)
    cost = calculate_content_tool_cost(usage.input_tokens, usage.output_tokens)
    is_ar = response_should_be_arabic(state, user_message)
    debug_payload = None
    if debug and settings.ENABLE_DEBUG_RESPONSE:
        debug_payload = {
            "original_filename": original_filename,
            "extracted_chars": len(resume_text),
            "resume_json": resume_json,
            "raw": provider_result.content,
            "provider_trace": provider_result.trace_metadata(),
        }

    file_info = GeneratedFileInfo(
        file_id=file_id,
        filename=filename,
        download_url=f"/tasks/resume-builder/download/{file_id}",
    )

    return ResumeBuilderChatResponse(
        user_id=user_id,
        sub_tool_id=sub_tool_id,
        conversation_uuid=conversation_uuid,
        message="تم إنشاء السيرة الذاتية بنجاح." if is_ar else "Resume generated successfully.",
        state=state,
        results=[ContentToolResultItem(id=1, text=state.last_output or "Resume generated", title="Resume Preview", meta={"file_id": file_id, "download_url": file_info.download_url})],
        count=1,
        file=file_info,
        request_id=request_id,
        debug=debug_payload,
        usage=usage,
        cost=cost,
    )
